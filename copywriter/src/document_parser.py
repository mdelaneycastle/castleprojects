"""
Document Parser Module
Extracts text from DOCX, PDF, HTML, and DOC files for style analysis and RAG.
"""

import os
import json
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Optional
from docx import Document
import pdfplumber


def detect_document_type(filename: str) -> str:
    """
    Detect document type based on filename.
    Returns: 'press_release', 'bio', or 'collection_overview'
    """
    filename_lower = filename.lower()

    if 'press release' in filename_lower:
        return 'press_release'
    elif 'bio' in filename_lower:
        return 'bio'
    elif 'overview' in filename_lower or 'collection' in filename_lower:
        return 'collection_overview'
    else:
        return 'general'


def parse_docx(file_input) -> str:
    """
    Extract text from a DOCX file, preserving paragraph structure.
    Accepts either a file path (str) or bytes.
    """
    from io import BytesIO

    if isinstance(file_input, bytes):
        doc = Document(BytesIO(file_input))
    else:
        doc = Document(file_input)

    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    return '\n\n'.join(paragraphs)


def parse_pdf(file_input) -> str:
    """
    Extract text from a PDF file using pdfplumber.
    Accepts either a file path (str) or bytes.
    """
    from io import BytesIO

    text_parts = []

    if isinstance(file_input, bytes):
        pdf_file = BytesIO(file_input)
    else:
        pdf_file = file_input

    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text.strip())

    return '\n\n'.join(text_parts)


def parse_html(file_input) -> str:
    """
    Extract text from an HTML file, stripping tags and preserving structure.
    Accepts either a file path (str) or bytes.
    """
    from io import BytesIO

    if isinstance(file_input, bytes):
        html_content = file_input.decode('utf-8', errors='ignore')
    else:
        with open(file_input, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()

    # Remove script and style elements
    html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
    html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)

    # Replace block-level elements with newlines
    html_content = re.sub(r'</?(p|div|br|h[1-6]|article|section|blockquote|li)[^>]*>', '\n', html_content, flags=re.IGNORECASE)

    # Remove all remaining HTML tags
    text = re.sub(r'<[^>]+>', '', html_content)

    # Decode HTML entities
    import html
    text = html.unescape(text)

    # Clean up whitespace
    lines = [line.strip() for line in text.split('\n')]
    paragraphs = []
    current_para = []

    for line in lines:
        if line:
            current_para.append(line)
        elif current_para:
            paragraphs.append(' '.join(current_para))
            current_para = []

    if current_para:
        paragraphs.append(' '.join(current_para))

    return '\n\n'.join(paragraphs)


def parse_doc(file_input) -> str:
    """
    Extract text from an old-style .doc file using macOS textutil.
    Accepts either a file path (str) or bytes.
    """
    from io import BytesIO

    if isinstance(file_input, bytes):
        # Write bytes to a temporary file
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
            tmp.write(file_input)
            tmp_path = tmp.name

        try:
            result = subprocess.run(
                ['textutil', '-convert', 'txt', '-stdout', tmp_path],
                capture_output=True,
                text=True
            )
            return result.stdout.strip()
        finally:
            os.unlink(tmp_path)
    else:
        result = subprocess.run(
            ['textutil', '-convert', 'txt', '-stdout', file_input],
            capture_output=True,
            text=True
        )
        return result.stdout.strip()


def parse_document_bytes(filename: str, file_bytes: bytes) -> dict:
    """
    Parse a document from bytes (for Streamlit uploads / Supabase).

    Returns:
        dict with keys: filename, doc_type, full_text, paragraphs, word_count
    """
    extension = Path(filename).suffix.lower()

    if extension == '.docx':
        full_text = parse_docx(file_bytes)
    elif extension == '.pdf':
        full_text = parse_pdf(file_bytes)
    elif extension in ('.html', '.htm'):
        full_text = parse_html(file_bytes)
    elif extension == '.doc':
        full_text = parse_doc(file_bytes)
    elif extension == '.txt':
        full_text = file_bytes.decode('utf-8', errors='ignore')
    else:
        raise ValueError(f"Unsupported file type: {extension}")

    paragraphs = [p.strip() for p in full_text.split('\n\n') if p.strip()]

    return {
        'filename': filename,
        'doc_type': detect_document_type(filename),
        'full_text': full_text,
        'paragraphs': paragraphs,
        'word_count': len(full_text.split()),
    }


def parse_document(file_path: str) -> dict:
    """
    Parse a document and return structured data.

    Returns:
        dict with keys:
            - filename: str
            - file_path: str
            - doc_type: str ('press_release', 'bio', 'collection_overview', 'general')
            - full_text: str
            - paragraphs: list[str]
            - word_count: int
    """
    path = Path(file_path)
    filename = path.name
    extension = path.suffix.lower()

    # Extract text based on file type
    if extension == '.docx':
        full_text = parse_docx(file_path)
    elif extension == '.pdf':
        full_text = parse_pdf(file_path)
    elif extension == '.html' or extension == '.htm':
        full_text = parse_html(file_path)
    elif extension == '.doc':
        full_text = parse_doc(file_path)
    else:
        raise ValueError(f"Unsupported file type: {extension}")

    # Split into paragraphs
    paragraphs = [p.strip() for p in full_text.split('\n\n') if p.strip()]

    # Count words
    word_count = len(full_text.split())

    return {
        'filename': filename,
        'file_path': str(path.absolute()),
        'doc_type': detect_document_type(filename),
        'full_text': full_text,
        'paragraphs': paragraphs,
        'word_count': word_count
    }


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """
    Split text into overlapping chunks for embedding.

    Args:
        text: The text to chunk
        chunk_size: Target size of each chunk in characters
        overlap: Number of characters to overlap between chunks

    Returns:
        List of text chunks
    """
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size

        # Try to break at a sentence or paragraph boundary
        if end < len(text):
            # Look for paragraph break first
            para_break = text.rfind('\n\n', start, end)
            if para_break > start + chunk_size // 2:
                end = para_break
            else:
                # Look for sentence break
                sentence_break = max(
                    text.rfind('. ', start, end),
                    text.rfind('! ', start, end),
                    text.rfind('? ', start, end)
                )
                if sentence_break > start + chunk_size // 2:
                    end = sentence_break + 1

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        start = end - overlap if end < len(text) else len(text)

    return chunks


def load_all_documents(directory: str) -> list[dict]:
    """
    Load and parse all supported documents from a directory.

    Args:
        directory: Path to directory containing documents

    Returns:
        List of parsed document dictionaries
    """
    documents = []
    supported_extensions = {'.docx', '.pdf', '.html', '.htm', '.doc'}

    for file_path in Path(directory).iterdir():
        if file_path.suffix.lower() in supported_extensions:
            try:
                doc = parse_document(str(file_path))
                documents.append(doc)
                print(f"Parsed: {file_path.name} ({doc['word_count']} words)")
            except Exception as e:
                print(f"Error parsing {file_path.name}: {e}")

    return documents


def save_processed_documents(documents: list[dict], output_dir: str) -> str:
    """
    Save processed documents to JSON for caching.

    Returns:
        Path to the saved file
    """
    output_path = Path(output_dir) / 'processed_documents.json'

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(documents, f, indent=2, ensure_ascii=False)

    return str(output_path)


def load_processed_documents(output_dir: str) -> Optional[list[dict]]:
    """
    Load previously processed documents from cache.

    Returns:
        List of documents if cache exists, None otherwise
    """
    cache_path = Path(output_dir) / 'processed_documents.json'

    if cache_path.exists():
        with open(cache_path, 'r', encoding='utf-8') as f:
            return json.load(f)

    return None


if __name__ == '__main__':
    # Test the parser
    import sys

    if len(sys.argv) > 1:
        test_dir = sys.argv[1]
    else:
        test_dir = '/Volumes/NewUser/CopyWriter'

    print(f"Loading documents from: {test_dir}")
    docs = load_all_documents(test_dir)
    print(f"\nLoaded {len(docs)} documents")

    for doc in docs:
        print(f"\n{doc['filename']}:")
        print(f"  Type: {doc['doc_type']}")
        print(f"  Words: {doc['word_count']}")
        print(f"  Paragraphs: {len(doc['paragraphs'])}")
